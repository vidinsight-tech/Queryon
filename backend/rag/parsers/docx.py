"""DOCX metin çıkarma (python-docx): paragraflar + tablolar."""
from __future__ import annotations

from pathlib import Path
from typing import Union

from docx import Document

from backend.core.exceptions import ExtractionError
from backend.rag.parsers.base import BaseParser, PathOrBytes
from backend.rag.types import ParsedContent


class DocxParser(BaseParser):
    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return ("docx",)

    def extract(self, source: PathOrBytes, filename_hint: str | None = None) -> ParsedContent:
        if isinstance(source, bytes):
            from io import BytesIO
            stream = BytesIO(source)
        else:
            stream = str(Path(source))
        try:
            doc = Document(stream)
        except Exception as e:
            raise ExtractionError(f"DOCX açılamadı: {e}", cause=e) from e
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts) if parts else ""
        return ParsedContent(text=text, source_type="docx", metadata={})
